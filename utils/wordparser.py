import os
import json
from docx import Document
import docx
import re
import ast
from docx.text.run import Run
from typing import Dict, List
import docxedit  # Make sure this is installed
import sys
import copy
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from utils import model_llm

class WordFileManager:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.sections = {}  # structured resume data
        self.formatting = {}  # global formatting metadata
        self.doc = None
        self.BULLETS_RE = re.compile(r"[•·●◦\-–—•]")
        self.STAR_FMT_RE = re.compile(r"[*_`]+")
        self.PATH_self_dir = os.path.join(os.path.dirname(__file__))
        self.OBJ_ModeLLM = model_llm.ModelLLM()
        self._get_known_sections(self.filepath)

    def _format_text(self, run: Run) -> str:
        text = run.text
        if run.bold and run.italic:
            return f"***{text}***"
        elif run.bold:
            return f"**{text}**"
        elif run.italic:
            return f"*{text}*"
        return text
    
    
    def docx_to_text(self, path: str) -> str:
        doc = docx.Document(path)
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    
    
    def _get_known_sections(self, docx_path: str) -> list[str]:
        """
        1. Reads a .docx resume.
        2. Sends the text to Gemini with a prompt that asks ONLY for section headers
        (including inferred 'SUMMARY' if no explicit title).
        3. Returns a Python list of headers in document order.
        """
        resume_text = self.docx_to_text(docx_path)

        system_prompt = {"role": "system", "content":  f"""
                        You will be given a resume document. Your task is to analyze its structure and return a list of section headers that are present in the resume.

                        Please return a JSON array of strings, where each string is a distinct section heading (e.g., "SUMMARY", "EDUCATION", "EXPERIENCE", etc.).
                        Also return the Title part of the resume which contains the NAME and other contact info as "TITLE_<all_info>".

                        ⚠️ Important:
                        - Return headers in the order they appear in the document.
                        - If the resume starts with a paragraph that seems to be a self-description or overview, include "SUMMARY" even if no such title is present.

                        Only return the array, no explanation.
                        """}
        input_prompt = {"role": "user", "content":  f"""
                        RESUME TEXT:
                        \"\"\"{resume_text}\"\"\"
                        """}
        
        prompt = [system_prompt, input_prompt]
        output = self.OBJ_ModeLLM.query(prompt)
        raw = output.strip()
        start = raw.find('[')
        end   = raw.rfind(']') + 1
        isolated_str = raw[start:end]
        header_list = ast.literal_eval(isolated_str)

        # Ensure it's a list of uppercase strings (deduplicated, order kept)
        self.known_sections = []
        seen    = set()
        for h in header_list:
            h_upper = h.strip().upper()
            if h_upper and h_upper not in seen:
                self.known_sections.append(h_upper)
                seen.add(h_upper)


    def normalize(self, s: str) -> str:
        s = s.strip()
        s = self.BULLETS_RE.sub(" ", s)
        s = self.STAR_FMT_RE.sub("", s)           # drop **, *
        s = re.sub(r"\s+", " ", s)           # collapse spaces/tabs
        s = re.sub(r"[.:;,\-–—]+$", "", s)   # strip trailing punct/dash
        return s.upper()

    def _get_section_name(self, para_text: str) -> str:
        
        cleaned = para_text.strip().upper()
        para_text_raw = self.normalize(para_text)
        if len(para_text_raw) > 30:
            return None
        for known_section in self.known_sections:
            if known_section == para_text_raw.strip():
                return cleaned
        return None

    def read(self):
        self.doc = Document(self.filepath)
        current_section = None

        for para in self.doc.paragraphs:
            para_text = "".join([self._format_text(run) for run in para.runs])
            if not para_text:
                continue

            section = self._get_section_name(para_text)
            if section:
                current_section = section
                self.sections[current_section] = {
                    "content": [],
                    "formatting": {
                        "font_size": para.runs[0].font.size.pt if para.runs[0].font.size else None,
                        "bold": para.runs[0].bold,
                        "italic": para.runs[0].italic
                    }
                }
                continue

            if current_section:
                self.sections[current_section]["content"].append(para_text)
        
        for section in self.known_sections:
            if "TITLE" in section:
                self.sections["TITLE"] = section.split("TITLE_")[1]

    def mark_updates_for_docxedit(self, change_suggestions: dict):
        """
        Combine Agent‑5 suggestions into self.sections so write() can
        do in‑place replacements with docxedit.

        change_suggestions: the dict loaded from agent5_gemini.json
        """
        changes = change_suggestions.get("resume_changes", [])
        for change in changes:
            section = change.get("section")
            if section == 'TITLE':
                continue
            # normalized_section = self.normalize(section)
            if section not in self.sections:
                print(f"\t\t [!] -WordParser- Section '{section}' not found, skipping.")
                continue

            content = self.sections[section]["content"]

            # ---------- replacements ----------
            if "replace" in change:
                orig = change["replace"].get("original", "").strip()
                upd  = change["replace"].get("updated", "").strip()
                matched = False
                for i, line in enumerate(content):
                    # Compare after stripping (Gemini returns trimmed text)
                    if line.strip() == orig:
                        # Preserve leading / trailing whitespace and tabs
                        prefix = line[:len(line) - len(line.lstrip())]
                        suffix = line[len(line.rstrip()):]
                        content[i] = f"{prefix}{orig}<updated>{upd}{suffix}"
                        matched = True
                        break
                if not matched:
                    print(f"\t\t [!] -WordParser- Could not find exact text in '{section}' for:\n{orig}\n")

            # ---------- deletions ----------
            if "delete" in change:
                to_del = change["delete"].strip()
                new_content = [ln for ln in content if ln.strip() != to_del]
                if len(new_content) == len(content):
                    print(f"\t\t [!] -WordParser- Could not find deletion target in '{section}':\n{to_del}\n")
                self.sections[section]["content"] = new_content

            # print(json.dumps(self.sections, indent=2, ensure_ascii=False))


    def replace_paragraph_runs(self, doc: Document, old_str_full: str, new_str_full: str) -> bool:
        """
        Replaces run-level text in a paragraph that matches `old_str_full` with aligned substrings
        from `new_str_full`, preserving formatting and handling edge cases with proportional distribution.
        """
        for paragraph in doc.paragraphs:
            # Cleaned match to handle whitespace/newlines (aligns with your read() and mark_updates_for_docxedit)
            cleaned_para = paragraph.text.replace('\n', ' ').strip()
            cleaned_old = old_str_full.replace('\n', ' ').strip()
            if cleaned_old == cleaned_para:  # Exact match on cleaned text (safer than 'in')
                runs = paragraph.runs
                if not runs:
                    # No runs: Add new text as a single run
                    paragraph.add_run(new_str_full)
                    return True

                # Get original run lengths (handle empty runs)
                old_lengths = [len(run.text) for run in runs]
                total_old = sum(old_lengths)
                if total_old == 0:
                    # All empty: Add to first or create one
                    if runs:
                        runs[0].text = new_str_full
                    else:
                        paragraph.add_run(new_str_full)
                    return True

                # Distribute new_str_full proportionally (fixes uneven lengths/short runs)
                new_parts = []
                pos = 0
                for length in old_lengths:
                    if length == 0:
                        new_parts.append('')  # Preserve empty runs (e.g., for spacing)
                        continue
                    # Proportional slice (scales to new length)
                    proportion = length / total_old
                    slice_len = int(len(new_str_full) * proportion)
                    slice_len = max(1, slice_len)  # Avoid zero-length for short runs
                    end = min(pos + slice_len, len(new_str_full))
                    new_parts.append(new_str_full[pos:end])
                    pos = end

                # Append any remainder to the last non-empty part (handles longer new text)
                if pos < len(new_str_full):
                    for i in range(len(new_parts) - 1, -1, -1):
                        if new_parts[i]:
                            new_parts[i] += new_str_full[pos:]
                            break
                    else:
                        new_parts[-1] += new_str_full[pos:]

                # Apply new parts to runs (preserves each run's formatting)
                for i, run in enumerate(runs):
                    run.text = new_parts[i]

                # Clean up empty runs (avoids bloat, but keeps formatting-only ones)
                for run in list(paragraph.runs):
                    if not run.text.strip() and not self.is_formatting_only_run(run):
                        run._element.getparent().remove(run._element)

                # If new text is much longer, add a new run for overflow (copies last run's formatting)
                if len(paragraph.text) > 1.5 * len(old_str_full):  # Adjustable threshold
                    last_run = paragraph.runs[-1] if paragraph.runs else None
                    if last_run and len(last_run.text) > 2 * old_lengths[-1]:
                        overflow_text = last_run.text[old_lengths[-1]:]
                        last_run.text = last_run.text[:old_lengths[-1]]
                        new_run = paragraph.add_run(overflow_text)
                        self.copy_formatting(last_run, new_run)

                return True
        return False
    

    def is_formatting_only_run(self, run):
        """Check if a run is empty but carries formatting (e.g., bold space)."""
        return not run.text.strip() and (run.bold or run.italic or run.underline or run.font.size)

    def copy_formatting(self, source_run, target_run):
        """Copy formatting from one run to another."""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        # Add more properties if needed (e.g., font.name, color)


    def new_write(self, output_doc_path):
        replacement_pairs = self.generate_replacement_pairs_from_docx()

        # Then apply them using docxedit:
        doc = Document(self.filepath)
        for old, new in replacement_pairs:
            docxedit.replace_string(doc, old_string=old, new_string=new)

        doc.save(output_doc_path)
    
    def write(self, output_path: str):
        """
        Uses docxedit to replace existing content without breaking formatting.
        Assumes original file still exists and edits in-place.
        """
        doc = Document(self.filepath)
        replacements_made = 0

        for section, data in self.sections.items():
            if section == "TITLE":
                continue
            for item in data["content"]:
                if "<updated>" in item:  # convention for marking updated lines
                    parts = item.split("<updated>")
                    if len(parts) == 2:
                        old, new = parts[0], parts[1]
                        try:
                            # replacement_pairs = self.generate_replacement_pairs_from_docx(doc, old, new)
                            # for replacement_pair in replacement_pairs:
                            if(self.replace_paragraph_runs(doc, old, new)):
                            # docxedit.replace_string(doc, old_string=old, new_string=new)
                                replacements_made += 1
                        except Exception as e:
                            print(f"\t\t [!] -WordParser- Could not replace:\n'{old}' ➝ '{new}'\nError: {e}")

        doc.save(output_path)
        # print(f"[+] Updated {replacements_made} item(s) in '{output_path}'")

    def export_json(self, output_path: str = None):
        """
        If output_path is provided, saves the JSON to file. Otherwise, returns the JSON as a Python dict.
        """
        if output_path:
            with open(output_path, 'w') as f:
                json.dump(self.sections, f, indent=4)
        else:
            return self.sections


if __name__ == "__main__":
    input_resume = r"D:\Personal\Projects\Python_Projects\projects\PEPPER\data\candidate_data\David.docx"
    # input_resume = r"data/candidate_data/default_resumes/Braunschweig Resume.docx"
    output_resume = r"D:\Personal\Projects\Python_Projects\projects\PEPPER\data\candidate_data\vedant_out.docx"

    # resume_json = r"D:\Personal\Projects\Python_Projects\projects\PEPPER\data\resume_data.json"
    resume_json_updated = r"D:\Personal\Projects\Python_Projects\projects\PEPPER\data\resume_data_updated.json"
    agent4_json = r"D:\Personal\Projects\Python_Projects\projects\PEPPER\data\agent5_gemini.json"
    output_json = r"D:\Personal\Projects\Python_Projects\projects\PEPPER\data\vedant_out.json"

    with open(agent4_json, 'r', encoding='utf-8') as f:
        resume_coach_feedback = json.load(f)

    OBJ = WordFileManager(input_resume)
    OBJ.read()
    OBJ.export_json(output_json)
    OBJ.mark_updates_for_docxedit(resume_coach_feedback)
    OBJ.export_json()
    # OBJ.write_new(resume_json_updated)
    OBJ.write(output_resume)
